"""
Run from project root:  python tools/generate_project_brief.py
Output:  AmanCH_ProjectBrief.docx  (saved to project root)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE  = RGBColor(0x1a, 0x3a, 0x6e)
GOLD  = RGBColor(0xc9, 0xa2, 0x27)
BLACK = RGBColor(0x0f, 0x1e, 0x3c)
GREY  = RGBColor(0x5a, 0x6a, 0x8a)

def set_font(run, size, bold=False, color=BLACK, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    run.font.name = "Calibri"

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    color = BLUE if level == 1 else BLACK
    set_font(run, 14 if level == 1 else 11, bold=True, color=color)
    return p

def add_body(doc, text, color=BLACK, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, 10.5, color=color, italic=italic)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, 10.5)
    return p

def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(label + " ")
    set_font(r1, 10.5, bold=True)
    r2 = p.add_run(value)
    set_font(r2, 10.5, color=GREY)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D8DFF0")
    pBdr.append(bottom)
    pPr.append(pBdr)

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Project Brief — Project Focus Weeks")
r.font.size  = Pt(20)
r.font.bold  = True
r.font.color.rgb = BLUE
r.font.name  = "Calibri"
pPr  = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot  = OxmlElement("w:bottom")
bot.set(qn("w:val"),   "single")
bot.set(qn("w:sz"),    "12")
bot.set(qn("w:space"), "4")
bot.set(qn("w:color"), "C9A227")
pBdr.append(bot)
pPr.append(pBdr)
doc.add_paragraph()

add_label_value(doc, "Client/Company Name:",
    "PowerCoders — Coding and IT training for refugees and asylum seekers in Switzerland")
add_label_value(doc, "Website:",
    "refugee-assistant-switzerland.vercel.app")

add_divider(doc)
add_heading(doc, "Project Title")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r1 = p.add_run("AmanCH — Refugee Assistant Switzerland")
set_font(r1, 12, bold=True, color=BLUE)
add_body(doc, "Aman = Peace · CH = Switzerland", color=GREY, italic=True)

add_divider(doc)
add_heading(doc, "Problem Definition")
add_body(doc,
    "Over 100,000 refugees and asylum seekers in Switzerland must navigate one of Europe's most "
    "complex asylum systems — 5 different permit types, 26 cantonal variations, critical legal "
    "deadlines (e.g. a 30-day appeal window that forfeits all rights if missed) — often without "
    "speaking German, French, or Italian, without legal support, and without access to clear, "
    "reliable information in their own language. No free, multilingual, permit-specific digital "
    "tool existed to guide them.")

add_divider(doc)
add_heading(doc, "Success Criteria")
add_body(doc, "A working web application where any refugee can:")
for b in [
    "Ask a question in their own language (Arabic, Tigrinya, Somali, Dari, Ukrainian, Turkish, etc.) and receive an accurate answer in that same language",
    "Select their permit type (N / F / B / C / S) and get answers specific to their legal situation",
    "Select their canton and get local migration office information",
    "See clickable source links to official Swiss websites (SEM, OSAR, ch.ch) for every answer",
    "Use voice input if they cannot type",
    "Access the app from any device, including mobile, with no account or login required",
]:
    add_bullet(doc, b)

add_divider(doc)
add_heading(doc, "Specific Requirements or Constraints")
for b in [
    "Must support 15+ languages — language detected automatically, no selection needed",
    "Must use only verified Swiss official sources — AI is forbidden from inventing URLs or fabricating legal facts (RAG pipeline with URL sanitisation)",
    "Must be completely free — no account, no paywall",
    "Must cover all 5 Swiss permit types: N, F, B, C, S",
    "Must work on mobile and desktop browsers",
    "Must use open-source tools wherever possible (SQLite, sentence-transformers, gTTS, FastAPI, React)",
    "API key required: Groq API (free tier) — for LLM + Whisper speech-to-text",
]:
    add_bullet(doc, b)

add_divider(doc)
add_heading(doc, "Skillsets Needed in the Team")
for b in [
    "Python — FastAPI, SQLite, NLP/embeddings (sentence-transformers), RAG pipeline",
    "JavaScript / React — Frontend development, component architecture, mobile-responsive UI",
    "UI/UX Design — Mobile-first design, accessibility, multilingual user experience",
    "API Integration — Groq API (LLaMA 3.3-70b LLM + Whisper STT), gTTS text-to-speech",
    "DevOps / Deployment — Vercel (React frontend), Render (FastAPI backend), Docker",
    "Domain Knowledge — Swiss asylum law, SEM/OSAR source curation (70+ official sources)",
]:
    add_bullet(doc, b)

add_divider(doc)
add_heading(doc, "Additional Resources for Participants")
for b in [
    "GitHub repository:         github.com/alitech19/Refugee-Assistant-Switzerland",
    "Live web app:              refugee-assistant-switzerland.vercel.app",
    "API (interactive docs):    amanch.onrender.com/docs",
    "Groq API (free key):       console.groq.com",
    "SEM — Swiss Migration Authority:  sem.admin.ch",
    "OSAR — Free legal aid for refugees:  osar.ch",
    "PowerCoders:               powercoders.org",
    "Contact: Ali Sulaiman — alisulaiman.it@gmail.com",
]:
    add_bullet(doc, b)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Built during the PowerCoders programme · powercoders.org")
set_font(r, 9, color=GREY, italic=True)

# Save to project root regardless of where script is run from
root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
out  = os.path.join(root, "AmanCH_ProjectBrief.docx")
doc.save(out)
print(f"Saved: {out}")
